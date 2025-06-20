import sys
import os
import re
import tiktoken
import tkinter as tk
from tkinter import filedialog, messagebox, PhotoImage, simpledialog
from tkinter import ttk
import threading
from datetime import datetime
import logging
import hashlib
import copy
import json
import base64
import time
import keyring
import getpass

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from openai import OpenAI

# Import de cryptography pour le chiffrement
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.backends import default_backend
from cryptography.fernet import Fernet

# Configuration du logging pour suivre l'exécution de l'application
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

"""
Ce fichier contient l'application principale MIIAApp. 
Elle permet de charger une clé API, de sélectionner un fichier HPR,
d'appeler l'API OpenAI pour analyser le contenu et finalement de générer un rapport au format Word.
"""

###############################################################################
# 1. Fonctions de chiffrement / déchiffrement
###############################################################################

def derive_key(password: str, salt: bytes) -> bytes:
    """Dérive une clé à partir d'un mot de passe et d'un sel via PBKDF2.
       Cette clé sera utilisée pour le chiffrement et le déchiffrement.
    """
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=100000,
        backend=default_backend()
    )
    return base64.urlsafe_b64encode(kdf.derive(password.encode()))

def encrypt_api_key(api_key: str, password: str, output_file: str = "config.enc") -> None:
    """Chiffre la clé API et enregistre le sel et la donnée chiffrée dans un fichier JSON.
       Ce fichier servira lors du chargement de l'application pour déchiffrer la clé.
    """
    salt = os.urandom(16)  # Génère un sel aléatoire
    key = derive_key(password, salt)
    fernet = Fernet(key)
    encrypted_api_key = fernet.encrypt(api_key.encode())
    config_data = {
        "salt": base64.urlsafe_b64encode(salt).decode(),
        "encrypted_api_key": encrypted_api_key.decode()
    }
    with open(output_file, "w") as f:
        json.dump(config_data, f)
    print(f"Clé API chiffrée enregistrée dans {output_file}.")

def get_resource_path(relative_path: str) -> str:
    """
    Retourne le chemin absolu vers un fichier.
    Si l'application est compilée (mode 'frozen'), le chemin est basé sur sys._MEIPASS,
    sinon on utilise le répertoire courant.
    """
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS  # type: ignore
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def load_api_key(password: str, config_file: str = None) -> str:
    """
    Déchiffre et retourne la clé API depuis le fichier de configuration chiffré.
    Utilise get_resource_path pour localiser config.enc.
    """
    if config_file is None:
        config_file = get_resource_path("config.enc")

    with open(config_file, "r") as f:
        config_data = json.load(f)
    salt = base64.urlsafe_b64decode(config_data["salt"])
    encrypted_api_key = config_data["encrypted_api_key"].encode()
    key = derive_key(password, salt)
    fernet = Fernet(key)
    api_key = fernet.decrypt(encrypted_api_key).decode()
    return api_key

###############################################################################
# 2. Fonctions pour stocker/récupérer le mot de passe via keyring
###############################################################################

def get_saved_password() -> str:
    """Récupère le mot de passe stocké dans le trousseau (keyring)."""
    return keyring.get_password("MIIA_app", "api_password")

def save_password(password: str):
    """Sauvegarde le mot de passe dans le trousseau (keyring)."""
    keyring.set_password("MIIA_app", "api_password", password)

def get_api_key_from_storage():
    """
    Récupère la clé API depuis le trousseau.
    Si le mot de passe n'est pas enregistré, il est demandé à l'utilisateur.
    Ensuite, la clé API est déchiffrée depuis le fichier de configuration.
    """
    root_temp = tk.Tk()
    root_temp.withdraw()  # On masque la fenêtre principale temporaire
    password = get_saved_password()
    if not password:
        password = simpledialog.askstring("Mot de passe", "Entrez le mot de passe pour déchiffrer la clé API :", show="*")
        if not password:
            messagebox.showerror("Erreur", "Le mot de passe est requis pour déchiffrer la clé API.")
            sys.exit()
        save_password(password)
    try:
        api_key = load_api_key(password)
    except Exception as e:
        messagebox.showerror("Erreur", f"Erreur lors du déchiffrement de la clé API : {e}")
        sys.exit()
    root_temp.destroy()
    return api_key

###############################################################################
# 3. Option de configuration (génération de config.enc)
###############################################################################

# Si l'argument "setup" est passé lors de l'exécution, on lance la configuration.
if len(sys.argv) > 1 and sys.argv[1] == "setup":
    # Demande à l'utilisateur la clé API et le mot de passe pour pouvoir la chiffrer
    api_key_input = input("Entrez la clé API OpenAI: ")
    password_input = getpass.getpass("Entrez le mot de passe pour chiffrer la clé API: ")
    encrypt_api_key(api_key_input, password_input)
    print("Configuration générée. Vous pouvez maintenant lancer l'application normalement.")
    sys.exit()

###############################################################################
# 4. Chargement de la clé API et initialisation du client
###############################################################################

API_KEY = get_api_key_from_storage()
client = OpenAI(api_key=API_KEY)
MODEL_NAME = "o3-mini"
response_cache = {}  # Dictionnaire pour mettre en cache les réponses de l'API

def get_desktop_path():
    """Retourne le chemin du bureau de l'utilisateur."""
    return os.path.join(os.path.expanduser("~"), "Desktop")

###############################################################################
# 5. Classe principale de l'application
###############################################################################

class MIIAApp:
    def __init__(self, root):
        """
        Initialise l'interface graphique et configure les styles.
        Définit également quelques variables pour gérer l'analyse.
        """
        self.root = root
        self.root.title("MIIA Bio Ard'Aisne")
        self.root.geometry("600x450")
        self.root.minsize(500, 350)
        self.root.configure(bg="#f0f8ff")

        # Configuration du style via ttk
        style = ttk.Style()
        style.theme_use("clam")

        primary_blue = "#4dd0e1"
        secondary_blue = "#b2ebf2"
        button_bg = "#00bcd4"
        button_fg = "white"
        header_color = "#ff7043"
        text_color = "#333333"

        # Stockage des couleurs dans l'instance pour pouvoir les réutiliser
        self.primary_blue = primary_blue
        self.secondary_blue = secondary_blue
        self.button_bg = button_bg
        self.button_fg = button_fg
        self.header_color = header_color
        self.text_color = text_color

        # Configuration des widgets ttk
        style.configure("TFrame", background="#f0f8ff")
        style.configure("TButton",
                        font=("Arial", 12, "bold"),
                        padding=10,
                        relief="raised",
                        background=self.button_bg,
                        foreground=self.button_fg,
                        borderwidth=2)
        style.map("TButton",
                  background=[("active", self.secondary_blue)],
                  foreground=[("active", self.button_fg)])
        style.configure("TLabel",
                        font=("Arial", 12),
                        background="#f0f8ff",
                        foreground=self.text_color)
        style.configure("Title.TLabel",
                        font=("Arial", 18, "bold", "italic"),
                        foreground=self.header_color,
                        background="#f0f8ff")
        style.configure("TProgressbar",
                        thickness=20,
                        troughcolor=self.secondary_blue,
                        background=self.primary_blue,
                        bordercolor=self.primary_blue)

        # Initialisation des variables pour le contenu analysé
        self.hpr_content = None
        self.chatgpt_answer = None
        self.analysis_done = False

        # Chargement du modèle de document
        self.template_doc = self.load_template()
        # Création des widgets de l'interface
        self.create_widgets()

    def process_file(self):
        """
        Fonction déclenchée lors du clic sur "Générer Rapport".
        Elle permet de sélectionner un fichier, de lire son contenu,
        d'appeler l'analyse via ChatGPT et de générer un rapport Word.
        """
        self.hpr_content = None
        self.chatgpt_answer = None
        self.analysis_done = False

        self.generate_btn.config(state="disabled")  # On désactive le bouton pour éviter les relances

        # Boîte de dialogue pour sélectionner le fichier HPR
        file_path = filedialog.askopenfilename(
            title="Sélectionnez un fichier HPR",
            filetypes=[("Fichiers HPR", "*.hpr;*.hl7;*.txt"), ("Tous fichiers", "*.*")]
        )
        if not file_path:
            self.generate_btn.config(state="normal")
            return

        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                self.hpr_content = f.read()
            logging.info("Fichier chargé.")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la lecture du fichier : {e}")
            self.generate_btn.config(state="normal")
            return

        # Configuration de la barre de progression
        self.progress.config(mode="indeterminate")
        self.progress.pack(pady=10)
        self.progress.start()
        self.status_label.config(text="Réception des données...")

        full_response_text = ""

        def analysis_task():
            """
            Tâche exécutée dans un thread pour appeler l'API OpenAI de manière asynchrone.
            Elle met à jour l'interface avec différents messages de statut.
            """
            nonlocal full_response_text
            try:
                status_messages = [
                    "Analyse en cours...",
                    "Traitement du texte...",
                    "Interprétation des résultats...",
                    "Génération du rapport...",
                    "Finalisation..."
                ]
                message_index = 0

                # Appel à l'API avec stream pour traiter les données au fur et à mesure
                for chunk in self.call_chatgpt_improved(self.hpr_content):
                    if chunk is not None:
                        print(chunk, end="", flush=True)
                        full_response_text += chunk
                        self.root.after(0, self.progress.start)
                        if message_index < len(status_messages):
                            self.root.after(0, self.update_status_label, status_messages[message_index])
                            message_index += 1

                self.chatgpt_answer = full_response_text

            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de l'analyse ChatGPT : {e}")
            finally:
                self.analysis_done = True
                # Arrêt de la barre de progression et remise à jour de l'interface
                self.root.after(0, self.progress.stop)
                self.root.after(0, self.progress.pack_forget)
                self.root.after(0, self.update_status_label, "")
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_filename = f"Interprétation_{timestamp}.docx"
                output_path = os.path.join(get_desktop_path(), safe_filename)
                self.root.after(0, self.generate_word_file, self.chatgpt_answer, output_path)
                self.root.after(0, lambda: self.generate_btn.config(state="normal"))

        threading.Thread(target=analysis_task, daemon=True).start()

    def create_widgets(self):
        """Crée et organise les widgets de l'interface graphique."""
        self.main_frame = ttk.Frame(self.root, padding=30, style="TFrame")
        self.main_frame.pack(fill="both", expand=True)

        title_label = ttk.Label(
            self.main_frame,
            text="Générateur de Rapport d'Analyse HPR",
            style="Title.TLabel"
        )
        title_label.pack(pady=20)

        info_label = ttk.Label(
            self.main_frame,
            text="Cliquez sur 'Générer Rapport' pour sélectionner un fichier HPR et lancer l'analyse.",
            wraplength=500,
            justify="center",
            foreground=self.text_color
        )
        info_label.pack(pady=10)

        self.generate_btn = ttk.Button(
            self.main_frame,
            text="Générer Rapport",
            command=self.process_file
        )
        self.generate_btn.pack(pady=20)

        self.progress = ttk.Progressbar(
            self.main_frame, orient="horizontal", mode="indeterminate", length=400
        )

        self.status_label = ttk.Label(
            self.main_frame,
            text="",
            justify="center",
            wraplength=400,
            font=("Arial", 10, "italic"),
            foreground=self.text_color
        )
        self.status_label.pack(pady=5)

        self.quit_btn = ttk.Button(
            self.main_frame,
            text="Quitter",
            command=self.root.quit
        )
        self.quit_btn.pack(pady=10)

    def load_template(self):
        """
        Charge le modèle de document Word (docx) qui servira de base pour le rapport.
        Si le fichier n'existe pas, l'application affiche une erreur et se ferme.
        """
        modele_path = get_resource_path("modele_entete.docx")
        if not os.path.exists(modele_path):
            messagebox.showerror("Erreur", f"Le fichier modèle {modele_path} est introuvable !")
            self.root.quit()
        try:
            doc = Document(modele_path)
            return doc
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors du chargement du modèle : {e}")
            self.root.quit()

    def load_prompt_from_file(self):
        """
        Charge le template du prompt depuis un fichier texte.
        Ce prompt sera utilisé pour formuler la demande à l'API OpenAI.
        """
        prompt_path = get_resource_path("prompt.txt")
        if not os.path.exists(prompt_path):
            messagebox.showerror("Erreur", f"Le fichier prompt {prompt_path} est introuvable !")
            self.root.quit()
        try:
            with open(prompt_path, "r", encoding="utf-8") as f:
                prompt_template = f.read()
                return prompt_template
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors du chargement du prompt : {e}")
            self.root.quit()

    def build_improved_prompt(self, hpr_content: str) -> str:
        """
        Construit le prompt final en insérant le contenu HPR dans le template précédemment chargé.
        """
        prompt_template = self.load_prompt_from_file()
        prompt = prompt_template.format(hpr_content=hpr_content)
        return prompt

    def call_chatgpt_improved(self, hpr_content: str):
        """
        Appelle l'API OpenAI en streaming afin de récupérer l'analyse du contenu.
        Le résultat est mis en cache pour éviter des appels redondants sur le même contenu.
        """
        prompt = self.build_improved_prompt(hpr_content)
        content_hash = hashlib.md5(hpr_content.encode('utf-8')).hexdigest()

        if content_hash in response_cache:
            logging.info("Réponse ChatGPT trouvée dans le cache.")
            yield response_cache[content_hash]
            return

        try:
            stream = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": prompt}],
                max_completion_tokens=4000,
                stream=True
            )
            full_answer = ""
            for chunk in stream:
                if chunk.choices[0].delta.content is not None:
                    chunk_content = chunk.choices[0].delta.content
                    full_answer += chunk_content
                    yield chunk_content
            response_cache[content_hash] = full_answer
            logging.info("Réponse ChatGPT mise en cache.")
        except Exception as e:
            logging.error("Erreur lors de l'appel à l'API OpenAI : %s", e)
            raise

    def generate_word_file(self, text: str, output_path: str):
        """
        Génère le rapport d'interprétation au format Word.
        Le rapport est basé sur le modèle chargé et enrichi avec le texte d'analyse.
        """
        # On travaille sur une copie du template pour ne pas modifier l'original
        doc = copy.deepcopy(self.template_doc)

        # --- Titre principal ---
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_title = title.add_run("Rapport d'Interprétation")
        run_title.bold = True
        run_title.font.size = Pt(16)
        doc.add_paragraph("\n")

        # --- Section "Informations Générales" ---
        general_info_heading = doc.add_paragraph()
        run_general_info_heading = general_info_heading.add_run("Informations Générales")
        run_general_info_heading.bold = True
        run_general_info_heading.font.size = Pt(14)
        doc.add_paragraph("\n")

        # Date et heure de génération du rapport
        date_generation = doc.add_paragraph()
        date_generation_run = date_generation.add_run("📅 Date de génération du rapport : ")
        date_generation_run.bold = True
        date_generation_run.font.size = Pt(12)
        date_generation.add_run(datetime.now().strftime("%d/%m/%Y %H:%M")).font.size = Pt(12)
        doc.add_paragraph("\n")

        # Avertissement
        avertissement = doc.add_paragraph()
        avertissement_run = avertissement.add_run("⚠ Ce document est une aide à l'interprétation et ne remplace pas un avis médical.")
        avertissement_run.bold = True
        avertissement_run.font.size = Pt(12)
        doc.add_paragraph("\n")

        # --- Section "Interprétation des Résultats Biologiques" ---
        interpretation_heading = doc.add_paragraph()
        run_interpretation_heading = interpretation_heading.add_run("Interprétation des Résultats Biologiques")
        run_interpretation_heading.bold = True
        run_interpretation_heading.font.size = Pt(14)
        doc.add_paragraph("\n")

        if text:
            # Pour chaque ligne générée, on ajoute des indicateurs en fonction du contenu
            for line in text.split("\n"):
                clean_line = line.strip()
                if clean_line:
                    if "Norma" in clean_line:
                        clean_line = "✅ " + clean_line
                    elif ("Légèrement" in clean_line or "Possible" in clean_line or 
                          "Surveillance" in clean_line or "Vérifier" in clean_line or 
                          "explorer" in clean_line):
                        clean_line = "⚠️ " + clean_line
                    elif "Indétectable" in clean_line or "Absence" in clean_line:
                        clean_line = "✔️ " + clean_line

                    para = doc.add_paragraph(clean_line)
                    para.runs[0].font.size = Pt(11)
        else:
            doc.add_paragraph("Aucune interprétation générée.", style="Normal")
        doc.add_paragraph("\n")

        try:
            doc.save(output_path)
            messagebox.showinfo("Succès", f"Le rapport a été généré avec succès :\n{output_path}")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la sauvegarde du rapport : {e}")

    def update_status_label(self, message):
        """Met à jour l'étiquette de statut dans l'interface graphique."""
        self.status_label.config(text=message)

def main():
    """Fonction principale qui démarre l'interface graphique."""
    root = tk.Tk()

    # Gestion de l'icône de l'application
    if getattr(sys, 'frozen', False):
        icon_path = os.path.join(sys._MEIPASS, "logo.ico")
    else:
        icon_path = "logo.ico"

    try:
        root.iconbitmap(icon_path)
    except Exception as e:
        print(f"Erreur lors du chargement de l'icône : {e}")

    app = MIIAApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()